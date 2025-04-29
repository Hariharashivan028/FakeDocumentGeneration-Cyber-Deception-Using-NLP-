#!/usr/bin/env python
# coding: utf-8

# In[2]:


from google.colab import drive
drive.mount('/content/drive')


# In[ ]:


get_ipython().system('pip install -U sentence-transformers')


# In[ ]:


get_ipython().system('pip install python-docx')


# In[ ]:


import pandas as pd
import numpy as np
import docx
import string
import re
import nltk
from nltk.corpus import stopwords
from nltk.tokenize import word_tokenize
import math


# In[ ]:


get_ipython().system('apt-get install -y xvfb # Install X Virtual Frame Buffer')
import os
os.system('Xvfb :1 -screen 0 1600x1200x16  &')    # create virtual display with size 1600x1200 and 16 bit color. Color can be changed to 24 or 8
os.environ['DISPLAY']=':1.0'    # tell X clients to use our virtual DISPLAY :1.0


# In[ ]:


from sentence_transformers import SentenceTransformer
model = SentenceTransformer('bert-base-nli-mean-tokens')
from sklearn.metrics.pairwise import cosine_similarity


# In[ ]:


nltk.download('stopwords')
nltk.download('punkt')
nltk.download('averaged_perceptron_tagger')
nltk.download('wordnet')
from nltk.stem import WordNetLemmatizer
from nltk.corpus import stopwords
nltk.download('omw-1.4')


# **Global Functions**

# In[ ]:


# remove stopwords
# remove unnecessary symbols i.e ? , . , ", .55 etc
def getCleanText(text) :

    # remove punctuations such as ‘!"#$%&'()*+,-./:;?@[\]^_`{|}~’
    exclude = set(string.punctuation)
    inclcude = ['?','.']
    newText = ''.join(ch for ch in text if ch not in exclude or ch in inclcude )
    # convert text to lowercase
    #newText=  newText.lower()
    # remove \n and \t ,\ ,digits
    # remove digits from text
    newText = re.sub("^\d+\s|\s\d+\s|\s\d+$", " ", newText)
    newText = re.sub('\n|\t|”|‘|’|“', '', newText)
    #words = [word for word in newText.split() if word.lower() not in stopwords.words('english')]
    #clean_text = " ".join(words)*-+

    return newText

## get part of speech of text ( Give cleaned text as parameter)
def get_pos(text):
      text1=nltk.word_tokenize(text)
      tags=nltk.pos_tag(text1)
      #print(tags)
      return tags

# get noun chunks  ( Adj*Noun )
# This function will return a chunk_tree
def get_chunks(tags):
    #NN noun, singular 'desk'
    #NNS noun plural 'desks'
    #NNP proper noun, singular 'Harrison'
    #NNPS proper noun, plural 'Americans'


    grammar = r"""
    NP: {<JJ|JJS|JJR>*<NN|NNS|NNP|NNPS>+}
    """
    cp = nltk.RegexpParser(grammar)
    chunk_tree=cp.parse(tags)
    return chunk_tree


def get_concept_from_chunks(chunk_tree,concept_list):
  if(chunk_tree.label()=="NP"):
      chunk=""
      for i in chunk_tree.leaves():
        if(chunk==""):
          chunk =i[0]
          continue
        chunk=chunk+" "+i[0]
      concept_list.append(chunk)
  for subtree in chunk_tree:
      if type(subtree) == nltk.tree.Tree:
          get_concept_from_chunks(subtree,concept_list)


def getConcept(text):
  tags=get_pos(text)
  chunks_tree=get_chunks(tags)
  concept_list=[]
  get_concept_from_chunks(chunks_tree,concept_list)
  return concept_list

# get frequency of each concept in a para f(c,pk)
# it will return dictionary of concept with frq

def getConFrq(text):
  concept_list=getConcept(text)
  con_frq={}
  for concept in concept_list:
    con_frq[concept]=con_frq.get(concept,0)+1
  return con_frq

# This function will return all the paragraphs in the document
def readtxtPara(filepath):
    doc = docx.Document(filepath)
    para_list = []
    for para in doc.paragraphs:
      #if paragraph contains less than 15 words then discard it
      if(len(para.text.split())>15):
        para_list.append(getCleanText(para.text))
    return para_list

def readtxtParaOriginal(filepath):
    doc = docx.Document(filepath)
    para_list = []
    for para in doc.paragraphs:
        para_list.append(para.text)
    return para_list

#This function will return cosine Similarity between two sentence

def getConsineSimilarity(sentences):
      sentence_embeddings = model.encode(sentences)
      return cosine_similarity(
                    [sentence_embeddings[0]],
                    sentence_embeddings[1:]
                                            )[0][0]


# In[ ]:


class Matrics:

    para_list=[] #list of paragraphs (cleaned) in document
    para_concept_frq_list=[] # list of dictionary  ( concept frq for a paragraph )
    all_concept=[] # list of all concept
    concept_graph=[]
    related_concept_graph=[]
    sequentiality=-10000000
    connectivity=-100000000
    dispersion =-100000000

    # parameterized constructor
    def __init__(self,paragraphs):
        #self.filepath=filepath
        self.para_list=paragraphs
        self.para_concept_frq_list=self.getParaConcepFrqtList()
        self.all_concept=self.getAllConcept()
        self.concept_graph=self.getConceptGraph()
        self.related_concept_graph=self.getRelatedGraph()
        self.sequentiality = self.getSequentiality()
        self.connectivity = self.getConnectivity()
        self.dispersion=self.getDispersion()

    # This function will return all the paragraphs in the document
    def readtxtPara(self):
        doc = docx.Document(self.filepath)
        para_list = []
        for para in doc.paragraphs:
           #if paragraph contains less than 15 words then discard it
          if(len(para.text.split())>15):
            para_list.append(getCleanText(para.text))
        return para_list


    def getParaConcepFrqtList(self):

        para_concept_frq_list=[]

        for para in self.para_list:
          para_concept_frq_list.append(getConFrq(para))

        return para_concept_frq_list

    def getAllConcept(self):
        allConcept=[]
        for paraConcept in self.para_concept_frq_list:
          for concept  in paraConcept:
            allConcept.append(concept)
        allConcept=[*set(allConcept)]

        return allConcept


    # it will return dictionary of dictionar
    # { a:{b:3,c:4},b:{a:3,c:4}}
    # arguments : list of all unique concepts in the document, concept dictionary of each paragraph


    def getConceptGraph(self):

      conceptGraph={}
      for concept in self.all_concept:
          temdic={}
          for paraConcept in self.para_concept_frq_list:
            if(paraConcept.get(concept,0)!=0):
              for concept1 in paraConcept:
                if(concept!=concept1):
                  temdic[concept1]=temdic.get(concept1,0)+1

          conceptGraph[concept]=temdic
      return  conceptGraph


    def getRelatedGraph(self):
        conceptGraphW2={}

        for concept in self.concept_graph:
          temp={}
          for concept2 in self.concept_graph[concept]:
            if(self.concept_graph[concept][concept2]>1):
              temp[concept2]=self.concept_graph[concept][concept2]
          if(len(temp)>0):
                conceptGraphW2[concept]=temp

        return   conceptGraphW2


    def get_key_para(self, concept):

        key_para=0
        max_significance=0
        k=0

        # get key paragraph & max_significance
        for paraConFrq in self.para_concept_frq_list :
          tem=paraConFrq.get(concept,0)* max(len(self.related_concept_graph.get(concept,{})),1)
          if(tem>max_significance):
            max_significance=tem
            key_para=k
          k=k+1

        return key_para

    def getComprehensionBurden(self, concept):

        key_para=0
        max_significance=0
        k=0

        # get key paragraph & max_significance
        for paraConFrq in self.para_concept_frq_list :
          tem=paraConFrq.get(concept,0)* max(len(self.related_concept_graph.get(concept,{})),1)
          if(tem>max_significance):
            max_significance=tem
            key_para=k
          k=k+1

        # Get Total burden  comprehension

        return key_para*max_significance

    def getSequentiality(self):

      sequentiality=0
      for concept in self.all_concept :

        sequentiality=sequentiality-self.getComprehensionBurden(concept)

      return sequentiality


    # get connectivity

    def getConnectivity(self):
      final_connectivity=0
      for concept in self.concept_graph:
        connectivity=0
        for concept2 in self.concept_graph[concept]:
            connectivity = connectivity+self.concept_graph[concept][concept2]

        final_connectivity=final_connectivity+pow(connectivity,1.2)

      C=len(self.concept_graph)
      final_connectivity=final_connectivity/C
      return final_connectivity


    # get Dispersion

    def getDispersion(self):
      total_dispersion=0
      for paraConFrq in self.para_concept_frq_list:

        dispersion=0

        total_concept=0
        for concept in paraConFrq:
          total_concept=total_concept+paraConFrq[concept]

        for concept in paraConFrq:
          tem=paraConFrq.get(concept,0)
          if(tem>0):
            dispersion=dispersion-(tem/total_concept)*math.log2(tem/total_concept)

        total_dispersion=total_dispersion+dispersion


      return total_dispersion/len(self.para_concept_frq_list)




# In[ ]:


# deepfrom sdeep.special import j0

class EditOperation:

    matrics=""
    concept_key_para = {}

    def __init__(self , matrics):
      self.matrics=matrics
      self.get_concept_key_para()


# Get key para for all concepts
    def get_concept_key_para(self):

        for concept in self.matrics.all_concept:
            self.concept_key_para[concept]=self.matrics.get_key_para(concept)

# Get all the sentences which contians concept having comprehension burden (Ψ(c) > 0)
    def get_sent_with_burden(self):

        para_sentence = []
        result=set()

        for para in self.matrics.para_list :
          para_sentence.append(para.split('. '))


        for concept in self.matrics.all_concept:

          for i in range(0, self.concept_key_para[concept]):
            for j in range(0,len(para_sentence[i])):
                if(para_sentence[i][j].find(concept)!=-1):
                  result.add(para_sentence[i][j])


        return list(result)






    def addition(self, sentence):

      # On each paragraph insert sentence and than check Sequentiality(d'), and similirity after insertion
      result=[]
      final_para_sent_list=[]

      for para in self.matrics.para_list:
        final_para_sent_list.append(para.split('. '))

      for i  in range(0,len(final_para_sent_list)):

      # If same paragraph then skip it

        if(self.matrics.para_list[i].find(sentence)!=-1):
          continue

        temp_sentences=list(final_para_sent_list[i])

        n =len(temp_sentences)
        similirity=0
        for j in range(0,n-1):
           similirity=similirity+getConsineSimilarity([temp_sentences[j],temp_sentences[j+1]])

        for j in range(0, n-2):


          tem_similirity=similirity-getConsineSimilarity([temp_sentences[j],temp_sentences[j+1]])
          tem_similirity=tem_similirity+getConsineSimilarity([temp_sentences[j],sentence])+getConsineSimilarity([sentence,temp_sentences[j+1]])

          temp_para_sent_list=[item[:] for item in final_para_sent_list]

          temp_para_sent_list[i].insert(j,sentence)

          temp_para_list=[]

          for para_sent in temp_para_sent_list:
            temp_para_list.append('. '.join(para_sent))


          if(self.matrics.sequentiality>Matrics(temp_para_list).sequentiality and similirity<tem_similirity):  # Modification required threashold value for similirity
           similirity = tem_similirity
           result=[i,j]

      return result


    def deletion (self,concept,para_list, visited):
      matrics=Matrics(para_list)

      print(concept)

      para_list=matrics.para_list


      key_para=0
      max_significance=0
      k=0

      # get key paragraph & max_significance
      for i in range(0, len(matrics.para_concept_frq_list)) :
        if(visited[i]==1):
          continue
        paraConFrq=matrics.para_concept_frq_list[i]
        tem=paraConFrq.get(concept,0)* max(len(matrics.related_concept_graph.get(concept,{})),1)
        if(tem>max_significance):
          max_significance=tem
          key_para=k
        k=k+1

      if(visited[key_para]==1):
        return para_list

      print(key_para)


      para_sentence = []

      for i in range(0,len(matrics.para_list)) :
        para=matrics.para_list[i]
        para_sentence.append(para.split('. '))

      for sent in para_sentence[key_para]:
        print(sent)
      print("----------------------------------------------------------")



      print(len(para_sentence[key_para]))

      for j in range(0,len(para_sentence[key_para])):
        print(j)
        if(para_sentence[key_para][j].find(concept)!=-1):
          print("----------------------------------------------------------")

          tem_para_sentence=para_sentence
          tem_para_sentence[key_para].remove(tem_para_sentence[key_para][j])
          visited[key_para]=1

          print("---------")
          print(tem_para_sentence)
          print("---------")

          print(para_sentence)

          temp_para_list=[]
          for para_sent in para_sentence:
            temp_para_list.append(". ".join(para_sent))

          print(matrics.sequentiality,Matrics(temp_para_list).sequentiality)

          if(matrics.sequentiality>Matrics(temp_para_list).sequentiality):
            return temp_para_list


      return para_list



    def deletion2(self, concept,visited,para_list):


        key_para=self.concept_key_para[concept]
        if(visited[key_para]==1):
          return list(para_list)


        key_para_sent=para_list[key_para].split('. ')

        for j in range(0, len(key_para_sent)):

          if(key_para_sent[j].find(concept)!=-1):
            temp_para_sents = list(key_para_sent)
            temp_para_sents.remove(key_para_sent[j])

            temp_para_list=list(para_list)
            temp_para_list[key_para]='. '.join(temp_para_sents)
            if(self.matrics.sequentiality>Matrics(temp_para_list).sequentiality):
              visited[key_para]=1
              print(key_para,key_para_sent[j])
              return temp_para_list

        return list(para_list)


    def minimum_required_frq(self, concept):
      result=0
      for i in range(self.concept_key_para[concept]+1,len(self.matrics.para_list)):
        count=self.matrics.para_list[i].count(concept)
        if(count>result):
          result=count

      return result


# In[ ]:





# In[ ]:


for i in range(1,11):
  path ="/content/drive/MyDrive/Paper_Implementation/Dataset/Original/encyclopedia/doc"+str(i)+".docx"
  path2 ="/content/drive/MyDrive/Paper_Implementation/Dataset/Believable_fake/encyclopedia/doc"+str(i)+".docx"
  original =Matrics(readtxtPara(path))
  fake =Matrics(readtxtPara(path2))
  print( "Original:",i,"->"," sequentiality",original.sequentiality,"connectivity",original.connectivity,"dispersion", original.dispersion)
  print( "Fake:",i,"->"," sequentiality",fake.sequentiality,"connectivity",fake.connectivity,"dispersion", fake.dispersion)
  if(original.sequentiality>=fake.sequentiality and original.connectivity>=fake.connectivity and original.dispersion<=fake.dispersion):
    print("Correct")
  else:
    print("Incorrect")


  print("-------------------------------------------------------------")


# In[ ]:


for i in range(1,2):
  path ="/content/drive/MyDrive/Paper_Implementation/Dataset/Original/encyclopedia/doc"+str(i)+".docx"
  # path2 ="/content/drive/MyDrive/Paper_Implementation/Dataset/Believable_fake/encyclopedia/doc"+str(i)+".docx"
  original =Matrics(readtxtPara(path))
  # fake =Matrics(readtxtPara(path2))
  print(original.concept_graph)


# In[ ]:


for i in range(1,11):
  path ="/content/drive/MyDrive/Paper_Implementation/Dataset/Original/CACM/doc"+str(i)+".docx"
  path2 ="/content/drive/MyDrive/Paper_Implementation/Dataset/Believable_fake/CACM/doc"+str(i)+".docx"
  original =Matrics(readtxtPara(path))
  fake =Matrics(readtxtPara(path2))
  print( "Original:",i,"->"," sequentiality",original.sequentiality,"connectivity",original.connectivity,"dispersion", original.dispersion)
  print( "Fake:",i,"->"," sequentiality",fake.sequentiality,"connectivity",fake.connectivity,"dispersion", fake.dispersion)
  if(original.sequentiality>=fake.sequentiality and original.connectivity>=fake.connectivity and original.dispersion<=fake.dispersion):
    print("Correct")
  else:
    print("Incorrect")


  print("-------------------------------------------------------------")


# In[ ]:


for i in range(1,11):
  path ="/content/drive/MyDrive/Paper_Implementation/Dataset/Original/Wikipedia/doc"+str(i)+".docx"
  path2 ="/content/drive/MyDrive/Paper_Implementation/Dataset/Believable_fake/Wikipedia/doc"+str(i)+".docx"
  original =Matrics(readtxtPara(path))
  fake =Matrics(readtxtPara(path2))
  print( "Original:",i,"->"," sequentiality",original.sequentiality,"connectivity",original.connectivity,"dispersion", original.dispersion)
  print( "Fake:",i,"->"," sequentiality",fake.sequentiality,"connectivity",fake.connectivity,"dispersion", fake.dispersion)
  if(original.sequentiality>=fake.sequentiality and original.connectivity>=fake.connectivity and original.dispersion<=fake.dispersion):
    print("Correct")
  else:
    print("Incorrect")


  print("-------------------------------------------------------------")

